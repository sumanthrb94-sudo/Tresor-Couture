#!/usr/bin/env python3
"""Generate a Word document from docs/LAUNCH-MANUAL.md."""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import parse_xml

ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / 'docs' / 'LAUNCH-MANUAL.md'
DOCX_PATH = ROOT / 'Tresor-Couture-Launch-Manual.docx'


def set_cell_border(cell, **kwargs):
    """Helper to draw table cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = parse_xml(r'<w:{} xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'.format(edge))
                tcBorders.append(element)
            for key, value in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(value))


def add_styled_paragraph(doc, text, style=None, bold=False, italic=False, color=None):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def parse_inline_formatting(paragraph, text):
    """Split text on **bold** and apply runs."""
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            paragraph.add_run(part[2:-2]).bold = True
        else:
            paragraph.add_run(part)


def add_markdown_table(doc, lines):
    """Parse a markdown table from lines and add to doc."""
    header = [c.strip() for c in lines[0].split('|') if c.strip()]
    rows = []
    for line in lines[2:]:
        if not line.strip():
            continue
        rows.append([c.strip() for c in line.split('|') if c.strip()])

    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = h
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            if i < len(cells):
                cells[i].text = val


def main():
    if not MD_PATH.exists():
        print(f"Markdown file not found: {MD_PATH}")
        sys.exit(1)

    md = MD_PATH.read_text(encoding='utf-8')
    lines = md.splitlines()

    doc = Document()

    # Title
    title = doc.add_heading('Tresor Couture — Production Launch Manual', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    in_code = False
    code_lines = []
    in_table = False
    table_lines = []

    i = 0
    while i < len(lines):
        line = lines[i]

        if line.strip().startswith('```'):
            if in_code:
                # flush code block
                p = doc.add_paragraph()
                p.style = 'Quote'
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        if '|' in line and not line.startswith('#') and not line.startswith('-') and not line.startswith('*'):
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        else:
            if in_table:
                add_markdown_table(doc, table_lines)
                in_table = False
                table_lines = []

        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
            i += 1
            continue
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
            i += 1
            continue
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
            i += 1
            continue
        if line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            doc.add_paragraph('_' * 60)
            i += 1
            continue

        # Bullet list
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = line.strip()[2:]
            parse_inline_formatting(p, text)
            i += 1
            continue

        # Numbered list (simple)
        m = re.match(r'^(\d+)\.\s+(.*)$', line.strip())
        if m:
            p = doc.add_paragraph(style='List Number')
            parse_inline_formatting(p, m.group(2))
            i += 1
            continue

        # Blockquote / note
        if line.strip().startswith('> '):
            p = doc.add_paragraph(style='Intense Quote')
            parse_inline_formatting(p, line.strip()[2:])
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        parse_inline_formatting(p, line)
        i += 1

    # Flush any remaining table
    if in_table:
        add_markdown_table(doc, table_lines)

    doc.save(DOCX_PATH)
    print(f"Generated: {DOCX_PATH}")


if __name__ == '__main__':
    main()
